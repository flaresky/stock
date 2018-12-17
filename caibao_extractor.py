# -*- coding: utf-8 -*-
import sys
import pprint
import json
import traceback
from docx import Document
import docx.document
from docx.oxml.table import *
from docx.oxml.text.run import *
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from constants import *
import mysqllib
import utils
import settings

stock_code = settings.extract_stock_code
year = 2015
# doc_name = u'E:\\Dropbox\\投资理财\\600104上汽集团\\上汽集团：%d年年度报告.docx'%(year)
# doc_name = u'E:\\Dropbox\\投资理财\\601318中国平安\\中国平安：%d年年度报告.docx'%(year)
# doc_name = u'E:\\Dropbox\\投资理财\\000858五粮液\\五粮液：%d年年度报告.docx'%(year)
# wordDoc = Document('/Users/flaresky/Downloads/2015.docx')
# wordDoc = Document('/Users/flaresky/Downloads/2013.docx')

def get_doc_name(stock_code, year):
#     stock_name = utils.get_stock_name_by_code(stock_code)
    stock_name = settings.extract_stock_name
    stock_file = u'C:\\Users\\droidhen200\\Dropbox\\投资理财\\%s%s\\%s：%d年年度报告.docx'%(stock_code,stock_name,stock_name,year)
#     print stock_file
    return stock_file
#     return u'E:\\Dropbox\\投资理财\\000858五粮液\\五粮液：%d年年度报告.docx'%(year)
#     return u'E:\\Dropbox\\投资理财\\000858五粮液\\五粮液：2011年年度报告2.docx'
#     return u'E:\\Dropbox\\投资理财\\600104上汽集团\\上汽集团：%d年年度报告.docx'%(year)
#     return u'/Users/flaresky/Dropbox/投资理财/600104上汽集团/上汽集团：%d年年度报告.docx'%(year)

def get_number(ele):
    num = get_text(ele, '0').replace(',', '')
    if num == '-':
        num = 0
    elif num[0] == '(' and num[-1] == ')':
        num = '-' + num[1:-1]
    return float(num)

# def get_text(ele, default=''):
#     for text in ele.itertext():
#         return text
#     return default
def get_text(ele, default=''):
    ret = None
    last_text = None
    for text in ele.itertext():
        if text == last_text: #会有重复的text，不知道为什么，过滤掉就没问题
            continue
        else:
            ret = text if ret is None else ret + text
            last_text = text
    if ret is not None:
        ret = ret.replace(' ', '')
    return default if ret is None else ret

def search_field(tb, field):
    mlen = 0
    mfield = None
    for f,show_f in TABLE_FIELDS[tb]:
        if field.find(f) >= 0:
            if len(f) > mlen:
                mlen = len(f)
                mfield = f if show_f is None else show_f
    if mfield is not None:
        return mfield
#     for f in TABLE_FIELDS[tb]:
#         if f.find(field) == 0:
#             return f
    return None

def save_caibao(tb, value_table):
    for year, vt in value_table.items():
        try:
            keys = ['code', 'year']
            keys.append(tb)
            keys = ["`"+f+"`" for f in keys]
            vals = [stock_code, year, json.dumps(vt, ensure_ascii=False)]
            vals = ["'"+unicode(f)+"'" for f in vals]
            updates = [keys[i]+"="+vals[i] for i in range(0, len(keys))]
        except:
            traceback.print_exc()
            return
        
        sql = "INSERT INTO stock_caibao (%s) VALUES (%s) ON DUPLICATE KEY UPDATE %s"%(', '.join(keys), ', '.join(vals), ', '.join(updates))
        mysqllib.run_sql(sql)
        print sql

def pretty_print(tb, value_table):
    res = '{\n'
    for year, vt in value_table.items():
        res += '\t%d : {\n'%(year)
        for k,nov in TABLE_FIELDS[tb]:
            if k in vt:
                res += '\t\t%s %s\n'%(k, vt[k])
        res += '\t}\n'
    res += '}\n'
    print res

def extract_table(stock_code, year, tb, print_title=False):
    value_table = {
                   year : {},
                   year-1 : {},
                   }
    is_in = False
    tr_len = 0
    pre_field = None
    pre_val_this = None
    pre_val_pre = None
    wordDoc = Document(get_doc_name(stock_code, year))
    for iter in wordDoc.element.getiterator():
        if isinstance(iter, CT_Text):
            t = iter.text.replace(' ', '')
            if t.find(TABLE_CONFIG[tb]['begin']['title']) >= 0 and len(t) <= TABLE_CONFIG[tb]['begin']['len']:
                print t
                is_in = True
            elif t.find(TABLE_CONFIG[tb]['end']['title']) >= 0 and len(t) <= TABLE_CONFIG[tb]['end']['len']:
                print t
                is_in = False
            pass
        elif isinstance(iter, CT_Tbl):
            if is_in:
                if tr_len == 0:
                    tr_len = iter.col_count
                    
                def table_iterator():
                    indexes = (0, 2, 3)
                    if tr_len == 3:
                        indexes = (0, 1, 2)
                    for tr in iter.tr_lst:
                        title = ''
                        val_this = 0
                        val_pre = 0
                        try:
                            for tc in tr.tc_lst:
                                if tc.left == indexes[0]:
                                    title = get_text(tc)
                                elif tc.left == indexes[1]:
                                    val_this = get_number(tc)
                                elif tc.left == indexes[2]:
                                    val_pre = get_number(tc)
                        except:
                            continue
                        yield (title, val_this, val_pre)
                    if tr_len >= 6:
                        indexes = (4, 6, 7)
                        if tr_len == 6:
                            indexes = (3, 4, 5)
                        for tr in iter.tr_lst:
                            title = ''
                            val_this = 0
                            val_pre = 0
                            try:
                                for tc in tr.tc_lst:
                                    if tc.left == indexes[0]:
                                        title = get_text(tc)
                                    elif tc.left == indexes[1]:
                                        val_this = get_number(tc)
                                    elif tc.left == indexes[2]:
                                        val_pre = get_number(tc)
                            except:
                                continue
                            yield (title, val_this, val_pre)
                
                for title, val_this, val_pre in table_iterator():
                    if not print_title:
                        field_name = search_field(tb, title)
                        if field_name is not None:
                            if val_this <> 0:
                                value_table[year][field_name] = val_this
                            if val_pre <> 0:
                                value_table[year-1][field_name] = val_pre
                            pre_field = None
                        else:
                            if pre_field is None:
                                pre_field = title
                                pre_val_this = val_this
                                pre_val_pre = val_pre
                            else:
                                pre_field += title
                                cat_field = search_field(tb, pre_field)
                                if cat_field is not None:
                                    if val_this <> 0:
                                        value_table[year][cat_field] = val_this
                                    elif pre_val_this <> 0:
                                        value_table[year][cat_field] = pre_val_this
                                    if val_pre <> 0:
                                        value_table[year-1][cat_field] = val_pre
                                    elif pre_val_pre <> 0:
                                        value_table[year-1][cat_field] = pre_val_pre
                                    pre_field = None
                    else:
                        print "u'%s',"%(title)
    pretty_print(tb, value_table)
    save_caibao(tb, value_table)

# for year in [2015, 2013, 2011]:
for year in [2013]:
#     extract_table(stock_code, year, 'zcfzb')
    extract_table(stock_code, year, 'lrb')
#     extract_table(stock_code, year, 'xjllb')
